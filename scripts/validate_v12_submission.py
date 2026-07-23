#!/usr/bin/env python3
"""Cross-validate every active v12 manuscript, figure, data and release claim."""
from __future__ import annotations

import hashlib
import json
import re
import subprocess
from pathlib import Path
from typing import Callable

import numpy as np
import pandas as pd
from docx import Document
from PIL import Image


HERE = Path(__file__).resolve().parent
SERVER = HERE.parent.parent
REPORTS = HERE / "reports"
FIGURES = HERE / "publication_figures_v12"
SOURCE = HERE / "v12_figure_source_data"
SUPP = HERE / "CoupledMD_Supplementary_Data_v12"
MANUSCRIPT = HERE / "v12_manuscript.md"
SI = HERE / "v12_si.md"
COHORT = SERVER / "data/release_cohort_v9_final207.csv"

tests: list[dict[str, str]] = []


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(8 * 1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def check(identifier: str, description: str, function: Callable[[], None]) -> None:
    try:
        function()
    except Exception as exc:
        tests.append(
            {
                "test_id": identifier,
                "description": description,
                "status": "FAIL",
                "detail": f"{type(exc).__name__}: {exc}",
            }
        )
    else:
        tests.append(
            {
                "test_id": identifier,
                "description": description,
                "status": "PASS",
                "detail": "",
            }
        )


def assert_equal(actual, expected) -> None:
    assert actual == expected, f"expected {expected!r}, observed {actual!r}"


def assert_file(path: Path) -> None:
    assert path.is_file() and path.stat().st_size > 0, path


def require(condition: bool, message: str = "requirement not satisfied") -> None:
    assert bool(condition), message


def markdown_checks() -> None:
    manuscript = MANUSCRIPT.read_text(encoding="utf-8")
    si = SI.read_text(encoding="utf-8")
    title = manuscript.splitlines()[0].removeprefix("# ")
    abstract = manuscript.split("## Abstract\n\n", 1)[1].split("\n\n## ", 1)[0]
    check(
        "TXT001",
        "title format",
        lambda: (
            assert_equal(len(title), 78),
            assert_equal(":" in title, False),
            assert_equal("CoupledMD" in title, False),
        ),
    )
    check(
        "TXT002",
        "abstract length and URL gate",
        lambda: (
            assert_equal(len(re.findall(r"\b[\wµ–-]+\b", abstract)), 164),
            assert_equal("http" in abstract, False),
        ),
    )
    required_sections = [
        "## Background & Summary",
        "## Methods",
        "## Data Records",
        "## Technical Validation",
        "## Usage Notes",
        "## Data Availability",
        "## Code Availability",
        "## Related manuscripts",
        "## Author Contributions",
        "## Competing Interests",
        "## References",
    ]
    check(
        "TXT003",
        "required manuscript sections",
        lambda: [
            require(section in manuscript, section)
            for section in required_sections
        ],
    )
    check(
        "TXT004",
        "no Results or Discussion section",
        lambda: (
            assert_equal("\n## Results" in manuscript, False),
            assert_equal("\n## Discussion" in manuscript, False),
        ),
    )
    check(
        "TXT005",
        "five main figures and no Figure 6",
        lambda: (
            [
                require(f"**Figure {number} |" in manuscript)
                for number in range(1, 6)
            ],
            assert_equal("**Figure 6 |" in manuscript, False),
        ),
    )
    required_truth = [
        "205 systems, giving 2,149",
        "205 detected + 2 explicit zero-pocket",
        "179 AMBER PMEMD systems",
        "310.4168 µs",
        "499.95 ns",
        "618 of 621",
        "classifies all 207",
        "computational repeats",
        "not biological replicates",
        "not the underlying full-system simulation archive",
        "reserved identifiers only",
    ]
    check(
        "TXT006",
        "required truth statements",
        lambda: [
            require(phrase in manuscript, phrase)
            for phrase in required_truth
        ],
    )
    prohibited = [
        r"\blargest\b",
        r"65 pocket clusters",
        r"C7[–-]Gq",
        r"Precog3D",
        r"\bBRET\b",
        r"\bepistasis\b",
        r"prospective challenge",
        r"model-selected mutation",
        r"archive is complete",
        r"full-system trajectory archive is available",
        r"all 621 primary trajectories are public",
    ]
    check(
        "TXT007",
        "prohibited novelty, resource and AI claims absent",
        lambda: [
            require(
                re.search(pattern, manuscript, flags=re.I) is None,
                pattern,
            )
            for pattern in prohibited
        ],
    )
    check(
        "TXT008",
        "SI ownership and release tables present",
        lambda: (
            require("Supplementary Table S2. Database-versus-resource claim and data overlap" in si),
            require("Supplementary Table S3. Current release and submission gate" in si),
            require("10.5281/zenodo.21395292" in si),
            require("private drafts" in si),
        ),
    )
    check(
        "TXT009",
        "active Markdown avoids ambiguous active counts",
        lambda: (
            require("182 Class A and 26 Class B" not in manuscript),
            require("42 Gq/11" not in manuscript),
            require("180 AMBER" not in manuscript),
            require("312.0 µs across" not in manuscript),
        ),
    )


def cohort_checks() -> set[str]:
    cohort = pd.read_csv(COHORT)
    cohort_ids = set(cohort.system_id)
    check(
        "COH001",
        "cohort row and ID count",
        lambda: (
            assert_equal(len(cohort), 207),
            assert_equal(cohort.system_id.nunique(), 207),
        ),
    )
    check(
        "COH002",
        "class counts",
        lambda: assert_equal(
            cohort.gpcr_class.value_counts().to_dict(),
            {"A": 181, "B": 26},
        ),
    )
    check(
        "COH003",
        "family counts",
        lambda: assert_equal(
            cohort.g_protein_family.value_counts().to_dict(),
            {"Gi": 95, "Gs": 65, "Gq": 41, "G12-13": 6},
        ),
    )
    check(
        "COH004",
        "class-family matrix",
        lambda: assert_equal(
            pd.crosstab(cohort.gpcr_class, cohort.g_protein_family)
            .to_dict(orient="index"),
            {
                "A": {"G12-13": 4, "Gi": 92, "Gq": 39, "Gs": 46},
                "B": {"G12-13": 2, "Gi": 3, "Gq": 2, "Gs": 19},
            },
        ),
    )
    check(
        "COH005",
        "receptor identifiers",
        lambda: (
            assert_equal(cohort.receptor_name.nunique(), 174),
            assert_equal(cohort.receptor_uniprot.nunique(), 173),
            assert_equal(
                cohort.loc[cohort.receptor_uniprot.isna(), "system_id"].tolist(),
                ["Gs_8HTI"],
            ),
        ),
    )
    check(
        "COH006",
        "nominal sampling",
        lambda: require(np.isclose(cohort.total_sampling_ns.sum(), 310500.0)),
    )
    return cohort_ids


def supplementary_checks(cohort_ids: set[str]) -> None:
    s1 = pd.read_csv(SUPP / "Supplementary_Data_S1_included_system_inventory.csv")
    s2 = pd.read_csv(SUPP / "Supplementary_Data_S2_release_boundary_exceptions.csv")
    s4 = pd.read_csv(SUPP / "Supplementary_Data_S4_reduced_release_manifest.csv")
    s5 = pd.read_csv(SUPP / "Supplementary_Data_S5_pocket_summaries_gpcrdb.csv")
    s6 = pd.read_csv(SUPP / "Supplementary_Data_S6_gateway_per_system.csv")
    s7 = pd.read_csv(SUPP / "Supplementary_Data_S7_source_simulation_replica_provenance.csv")
    s8 = pd.read_csv(SUPP / "Supplementary_Data_S8_reduced_release_qc.csv")
    s9 = pd.read_csv(SUPP / "Supplementary_Data_S9_api_access_audit.csv")
    check(
        "SUP001",
        "S1 exact cohort",
        lambda: (
            assert_equal(len(s1), 207),
            assert_equal(set(s1.system_id), cohort_ids),
        ),
    )
    check(
        "SUP002",
        "S2 boundary exceptions",
        lambda: (
            assert_equal(len(s2), 15),
            assert_equal(s2.release_status.value_counts().to_dict(), {"unresolved": 13, "excluded": 2}),
            assert_equal(set(s2.system_id) & cohort_ids, set()),
            assert_equal(set(s2.loc[s2.release_status.eq("excluded"), "system_id"]), {"Gq_7DWC", "Gq_7E9W"}),
        ),
    )
    check(
        "SUP003",
        "S4 expected file grid and no absolute archive paths",
        lambda: (
            assert_equal(len(s4), 207 * 3 * 2),
            assert_equal(set(s4.system_id), cohort_ids),
            assert_equal(set(s4.release_replica), {1, 2, 3}),
            assert_equal(set(s4.file_role), {"structure_pdb", "trajectory_xtc"}),
            require(not s4.archive_member_path.dropna().astype(str).str.startswith("/").any()),
        ),
    )
    detected = s5[s5.pocket_record_status.eq("available_detected_pocket")]
    zeros = s5[s5.pocket_record_status.eq("available_zero_pockets")]
    check(
        "SUP004",
        "S5 pocket totals and explicit zero records",
        lambda: (
            assert_equal(len(detected), 2149),
            assert_equal(detected.system_id.nunique(), 205),
            assert_equal(set(zeros.system_id), {"Gi_7YK6", "Gi_8YIC"}),
            assert_equal(s5.system_id.nunique(), 207),
        ),
    )
    check(
        "SUP005",
        "S6 gateway dimensions",
        lambda: (
            assert_equal(len(s6), 5796),
            assert_equal(s6.system_id.nunique(), 207),
            assert_equal(s6.interface.nunique(), 7),
            assert_equal(s6.metric.nunique(), 4),
            assert_equal(set(s6.system_id), cohort_ids),
        ),
    )
    check(
        "SUP006",
        "S7 replica provenance and sampling",
        lambda: (
            assert_equal(len(s7), 621),
            require(s7.groupby("system_id").replica_id.nunique().eq(3).all()),
            require(np.isclose(s7.nominal_duration_ns.sum() / 1000, 310.5)),
            require(np.isclose(s7.original_file_observed_duration_ns.sum() / 1000, 310.4168)),
            assert_equal(bool(s7.full_system_source_distributed.any()), False),
        ),
    )
    check(
        "SUP007",
        "S8 expected QC grid",
        lambda: (
            assert_equal(len(s8), 621),
            assert_equal(set(s8.system_id), cohort_ids),
            require(s8.groupby("system_id").release_replica.nunique().eq(3).all()),
        ),
    )
    check(
        "SUP008",
        "S9 API consensus denominator",
        lambda: require("ALL-group denominators observed: [207]" in "|".join(s9.availability_semantics.astype(str))),
    )
    check(
        "SUP009",
        "S3 and S10 are populated",
        lambda: (
            require(len(pd.read_csv(SUPP / "Supplementary_Data_S3_metadata_dictionary.csv")) > 50),
            require(len(pd.read_csv(SUPP / "Supplementary_Data_S10_code_environment_inventory.csv")) >= 10),
        ),
    )
    check(
        "SUP010",
        "Supplementary package audit and scope",
        lambda: (
            assert_equal(load_json(SUPP / "Supplementary_Data_package_audit.json")["scope"]["full_system_source_distributed"], False),
            require("not a full-system trajectory archive" in (SUPP / "README.txt").read_text(encoding="utf-8")),
        ),
    )


def load_json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def figure_checks(cohort_ids: set[str]) -> None:
    stems = [
        "figure1_cohort_scope",
        "figure2_records_reuse_boundary",
        "figure3_annotation_validation",
        "figure4_replica_technical_validation",
        "figure5_release_integrity_reuse",
        "supplementary_figure_s1_gateway_provenance",
        "supplementary_figure_s2_sampling_provenance",
    ]
    check(
        "FIG001",
        "all PDF/PNG figure pairs",
        lambda: [
            (assert_file(FIGURES / f"{stem}.pdf"), assert_file(FIGURES / f"{stem}.png"))
            for stem in stems
        ],
    )
    check(
        "FIG002",
        "PNG resolution",
        lambda: [
            assert_minimum_image(FIGURES / f"{stem}.png", 1800, 1400)
            for stem in stems
        ],
    )
    f1 = pd.read_csv(SOURCE / "Figure_1A_cohort_composition.csv")
    f3 = pd.read_csv(SOURCE / "Figure_3A_positive_control.csv")
    f4a = pd.read_csv(SOURCE / "Figure_4A_replica_validation.csv")
    f4b = pd.read_csv(SOURCE / "Figure_4B_reduced_record_qc.csv")
    f5 = pd.read_csv(SOURCE / "Figure_5A_release_status.csv")
    check(
        "FIG003",
        "Figure 1 source totals",
        lambda: assert_equal(int(f1.systems.sum()), 207),
    )
    check(
        "FIG004",
        "Figure 3 positive-control totals",
        lambda: (
            assert_equal(len(f3), 58),
            assert_equal(int(f3.ortho_recovered.sum()), 49),
        ),
    )
    check(
        "FIG005",
        "Figure 4 source totals",
        lambda: (
            assert_equal(len(f4a), 621),
            assert_equal(int(f4a.validation_status.eq("available").sum()), 618),
            assert_equal(len(f4b), 207),
            assert_equal(set(f4b.system_id), cohort_ids),
            assert_equal(f4b.status.value_counts().to_dict(), {"OK": 207}),
        ),
    )
    check(
        "FIG006",
        "Figure 5 status source does not claim publication",
        lambda: (
            assert_equal(len(f5), 3),
            assert_equal(bool(f5.publicly_resolvable.any()), False),
            require(set(f5.remote_state) <= {"unsubmitted", "not_verified"}),
        ),
    )
    manifest = pd.read_csv(SOURCE / "figure_source_data_manifest.csv")
    check(
        "FIG007",
        "figure-source manifest exact and hash-valid",
        lambda: validate_source_manifest(manifest),
    )
    check(
        "FIG008",
        "no active Figure 6 output",
        lambda: assert_equal(
            any("figure6" in path.name.lower() for path in FIGURES.iterdir()),
            False,
        ),
    )


def assert_minimum_image(path: Path, width: int, height: int) -> None:
    with Image.open(path) as image:
        assert image.width >= width and image.height >= height, (
            path,
            image.size,
        )


def validate_source_manifest(manifest: pd.DataFrame) -> None:
    expected = {
        path.name
        for path in SOURCE.glob("*.csv")
        if path.name != "figure_source_data_manifest.csv"
    }
    assert_equal(set(manifest.file), expected)
    assert_equal(len(manifest), 15)
    for row in manifest.itertuples(index=False):
        path = SOURCE / row.file
        assert_equal(sha256(path), row.sha256)
        assert_equal(len(pd.read_csv(path)), int(row.rows))


def api_and_release_checks() -> None:
    gateway = load_json(SERVER / "data/api/v1/consensus/gateways.json")
    check(
        "API001",
        "gateway cache final-207 denominators",
        lambda: assert_equal(
            {
                group: sorted(
                    {
                        int(row["n_systems"])
                        for row in gateway["records"]
                        if row["group"] == group
                    }
                )
                for group in ["ALL", "ClassA", "ClassB", "Gi", "Gs", "Gq", "G12-13"]
            },
            {
                "ALL": [207],
                "ClassA": [181],
                "ClassB": [26],
                "Gi": [95],
                "Gs": [65],
                "Gq": [41],
                "G12-13": [6],
            },
        ),
    )
    numerical = load_json(REPORTS / "v12_numerical_consistency_audit.json")
    check(
        "AUD001",
        "numerical audit exact cohort facts",
        lambda: (
            assert_equal(numerical["cohort"]["systems"], 207),
            assert_equal(numerical["replicas_and_sampling"]["nominal_replicas"], 621),
            require(np.isclose(numerical["replicas_and_sampling"]["nominal_sampling_us"], 310.5)),
            require(np.isclose(numerical["replicas_and_sampling"]["original_file_observed_sampling_us"], 310.4168)),
            assert_equal(numerical["pockets"]["systems_with_detected_pockets"], 205),
            assert_equal(numerical["pockets"]["positive_control_recovered"], 49),
        ),
    )
    remote = load_json(REPORTS / "v12_remote_release_verification.json")
    check(
        "AUD002",
        "remote release claims correctly withheld",
        lambda: (
            assert_equal(remote["summary"]["public_release_claim_allowed"], False),
            assert_equal(remote["summary"]["archive_completeness_claim_allowed"], False),
            assert_equal(remote["summary"]["all_three_public"], False),
        ),
    )
    check(
        "AUD003",
        "reference registry audit",
        lambda: assert_equal(load_json(REPORTS / "v12_reference_doi_audit.json")["status"], "PASS"),
    )
    check(
        "AUD004",
        "GitHub repository identity and access audit",
        lambda: assert_equal(
            load_json(REPORTS / "v12_repository_access_verification.json")[
                "status"
            ],
            "PASS",
        ),
    )
    check(
        "AUD005",
        "claim ledger and firewall reports",
        lambda: (
            require(len(pd.read_csv(REPORTS / "v12_manuscript_claim_evidence_ledger.csv")) >= 40),
            assert_file(REPORTS / "v12_database_resource_claim_firewall.md"),
            assert_file(REPORTS / "v12_revision_log.md"),
            assert_file(REPORTS / "v12_submission_readiness_checklist.json"),
        ),
    )


def document_checks() -> None:
    main_doc = HERE / "v12_manuscript_release_candidate.docx"
    si_doc = HERE / "v12_si_release_candidate.docx"
    main_pdf = HERE / "rendered_v12_manuscript/v12_manuscript_release_candidate.pdf"
    si_pdf = HERE / "rendered_v12_si/v12_si_release_candidate.pdf"
    check(
        "DOC001",
        "DOCX and PDF outputs exist",
        lambda: [assert_file(path) for path in [main_doc, si_doc, main_pdf, si_pdf]],
    )
    main = Document(main_doc)
    supplement = Document(si_doc)
    main_text = "\n".join(paragraph.text for paragraph in main.paragraphs)
    si_text = "\n".join(paragraph.text for paragraph in supplement.paragraphs)
    check(
        "DOC002",
        "DOCX embedded figure counts",
        lambda: (
            assert_equal(len(main.inline_shapes), 5),
            assert_equal(len(supplement.inline_shapes), 2),
        ),
    )
    check(
        "DOC003",
        "DOCX key text survives rendering source",
        lambda: (
            require("205 detected + 2 explicit zero-pocket" in main_text),
            require("restricted, unsubmitted private drafts" in main_text),
            require("Database-versus-resource claim and data overlap" in si_text),
        ),
    )
    check(
        "DOC004",
        "PDF page counts",
        lambda: (
            assert_equal(pdf_pages(main_pdf), 15),
            assert_equal(pdf_pages(si_pdf), 8),
        ),
    )
    visual = load_json(REPORTS / "v12_visual_qc_report.json")
    check(
        "DOC005",
        "visual QC passes",
        lambda: (
            assert_equal(visual["status"], "PASS"),
            assert_equal(visual["manual_visual_review"]["status"], "PASS"),
            assert_equal(visual["documents"]["manuscript"]["pages"], 15),
            assert_equal(visual["documents"]["supplementary_information"]["pages"], 8),
        ),
    )


def pdf_pages(path: Path) -> int:
    result = subprocess.run(
        ["pdfinfo", str(path)],
        check=True,
        capture_output=True,
        text=True,
    )
    match = re.search(r"^Pages:\s+(\d+)$", result.stdout, flags=re.M)
    assert match is not None
    return int(match.group(1))


def licence_checks() -> None:
    citation = (SERVER / "CITATION.cff").read_text(encoding="utf-8")
    check(
        "LIC001",
        "licence and citation metadata",
        lambda: (
            require("Jianxiang Huang" in (SERVER / "LICENSE").read_text(encoding="utf-8")),
            require("Jianxiang Huang" in (SERVER / "LICENSE-CODE").read_text(encoding="utf-8")),
            require("10.0000/placeholder" not in citation),
            require("https://github.com/HuangJianxiang-SJTU/coupledmd" in citation),
            require("https://github.com/HuangJianxiang-SJTU/coupledmd_database" in citation),
        ),
    )


def write_report() -> int:
    failed = [test for test in tests if test["status"] == "FAIL"]
    payload = {
        "schema_version": "1.0",
        "status": "PASS" if not failed else "FAIL",
        "summary": {
            "tests": len(tests),
            "passed": len(tests) - len(failed),
            "failed": len(failed),
            "submission_readiness_is_separate": True,
        },
        "tests": tests,
    }
    (REPORTS / "v12_submission_validation.json").write_text(
        json.dumps(payload, indent=2) + "\n",
        encoding="utf-8",
    )
    lines = [
        "# v12 automated submission-package validation",
        "",
        f"Status: **{payload['status']}**",
        "",
        f"Tests: {len(tests)}; passed: {len(tests) - len(failed)}; failed: {len(failed)}.",
        "",
        "This report validates internal artifact consistency. It does not override the separate submission-readiness BLOCKED/HOLD gates for data access.",
        "",
        "| Test | Description | Status | Detail |",
        "|---|---|---|---|---|",
    ]
    lines.extend(
        f"| {row['test_id']} | {row['description']} | {row['status']} | {row['detail']} |"
        for row in tests
    )
    (REPORTS / "v12_submission_validation.md").write_text(
        "\n".join(lines) + "\n",
        encoding="utf-8",
    )
    print(
        f"v12 validation: {len(tests) - len(failed)}/{len(tests)} PASS; "
        f"status={payload['status']}"
    )
    if failed:
        for item in failed:
            print(f"  FAIL {item['test_id']}: {item['detail']}")
    return 0 if not failed else 1


def main() -> int:
    markdown_checks()
    cohort_ids = cohort_checks()
    supplementary_checks(cohort_ids)
    figure_checks(cohort_ids)
    api_and_release_checks()
    document_checks()
    licence_checks()
    return write_report()


if __name__ == "__main__":
    raise SystemExit(main())
