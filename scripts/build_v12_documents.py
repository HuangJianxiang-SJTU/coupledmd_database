#!/usr/bin/env python3
"""Render auditable v12 Markdown sources to production DOCX files."""
from __future__ import annotations

import re
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


HERE = Path(__file__).resolve().parent
BUILDS = [
    (
        HERE / "v12_manuscript.md",
        HERE / "v12_manuscript_release_candidate.docx",
        True,
    ),
    (
        HERE / "v12_si.md",
        HERE / "v12_si_release_candidate.docx",
        False,
    ),
]


def set_cell_shading(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def set_repeat_header(row) -> None:
    properties = row._tr.get_or_add_trPr()
    header = properties.find(qn("w:tblHeader"))
    if header is None:
        header = OxmlElement("w:tblHeader")
        header.set(qn("w:val"), "true")
        properties.append(header)


def prevent_row_split(row) -> None:
    properties = row._tr.get_or_add_trPr()
    if properties.find(qn("w:cantSplit")) is None:
        properties.append(OxmlElement("w:cantSplit"))


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, end])


def configure_document(document: Document, line_numbers: bool) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.82)
    section.right_margin = Inches(0.82)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.35)
    add_page_number(section.footer.paragraphs[0])
    if line_numbers:
        properties = section._sectPr
        line_number = OxmlElement("w:lnNumType")
        line_number.set(qn("w:countBy"), "1")
        line_number.set(qn("w:restart"), "newPage")
        line_number.set(qn("w:distance"), "360")
        properties.append(line_number)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.widow_control = True

    for name, size, color in [
        ("Title", 17, "1E252B"),
        ("Heading 1", 13, "1E252B"),
        ("Heading 2", 11.5, "1E252B"),
        ("Heading 3", 10.5, "287C7E"),
    ]:
        style = document.styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(10 if name != "Title" else 0)
        style.paragraph_format.space_after = Pt(5)

    caption = document.styles["Caption"]
    caption.font.name = "Arial"
    caption.font.size = Pt(8.5)
    caption.font.color.rgb = RGBColor.from_string("30363B")
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    caption.paragraph_format.line_spacing = 1.05
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(7)
    caption.paragraph_format.keep_together = True

    if "Code Block" not in document.styles:
        code = document.styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code = document.styles["Code Block"]
    code.font.name = "Liberation Mono"
    code.font.size = Pt(7.5)
    code._element.rPr.rFonts.set(qn("w:eastAsia"), "Liberation Mono")
    code.paragraph_format.left_indent = Inches(0.18)
    code.paragraph_format.right_indent = Inches(0.12)
    code.paragraph_format.space_before = Pt(3)
    code.paragraph_format.space_after = Pt(5)
    code.paragraph_format.keep_together = True
    code_element = code._element.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "F3F5F6")
    code_element.append(shading)


INLINE = re.compile(
    r"(\*\*.+?\*\*|(?<!\*)\*[^*]+?\*(?!\*)|`[^`]+?`|\^[^^]+?\^)"
)


def add_inline(paragraph, text: str) -> None:
    text = text.replace(r"\*", "*")
    position = 0
    for match in INLINE.finditer(text):
        if match.start() > position:
            paragraph.add_run(text[position : match.start()])
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("*"):
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Liberation Mono"
            run.font.size = Pt(8.5)
            run.font.color.rgb = RGBColor.from_string("225E68")
        elif token.startswith("^"):
            run = paragraph.add_run(token[1:-1])
            run.font.superscript = True
            run.font.size = Pt(7.5)
        position = match.end()
    if position < len(text):
        paragraph.add_run(text[position:])


def add_heading(document: Document, text: str, level: int) -> None:
    if level == 1:
        paragraph = document.add_paragraph(style="Title")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        paragraph = document.add_heading(level=min(level - 1, 3))
    add_inline(paragraph, text)


def add_table(document: Document, lines: list[str]) -> None:
    parsed = []
    for line in lines:
        cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
        if all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells):
            continue
        parsed.append(cells)
    if not parsed:
        return
    width = len(parsed[0])
    if any(len(row) != width for row in parsed):
        raise ValueError(f"inconsistent Markdown table: {lines}")
    table = document.add_table(rows=len(parsed), cols=width)
    table.style = "Table Grid"
    table.autofit = True
    for row_index, values in enumerate(parsed):
        row = table.rows[row_index]
        prevent_row_split(row)
        if row_index == 0:
            set_repeat_header(row)
        for cell, value in zip(row.cells, values):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_index == 0:
                set_cell_shading(cell, "DCE8EA")
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            add_inline(paragraph, value)
            for run in paragraph.runs:
                run.font.size = Pt(7.7)
                if row_index == 0:
                    run.bold = True
    after = document.add_paragraph()
    after.paragraph_format.space_after = Pt(0)


def add_image(document: Document, markdown: str) -> None:
    match = re.fullmatch(r"!\[(.*?)\]\((.*?)\)", markdown.strip())
    if match is None:
        raise ValueError(markdown)
    path = HERE / match.group(2)
    if not path.is_file():
        raise FileNotFoundError(path)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.paragraph_format.space_after = Pt(1)
    paragraph.add_run().add_picture(str(path), width=Inches(6.75))


def add_paragraph(document: Document, text: str) -> None:
    is_caption = text.startswith("**Figure ") or text.startswith(
        "**Supplementary Figure "
    )
    style = "Caption" if is_caption else "Normal"
    paragraph = document.add_paragraph(style=style)
    if text.startswith(("Jianxiang Huang", "^1^", "^2^", "^3^", r"\*Correspondence")):
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.line_spacing = 1.0
    add_inline(paragraph, text)


def render(source: Path, target: Path, line_numbers: bool) -> None:
    document = Document()
    configure_document(document, line_numbers)
    lines = source.read_text(encoding="utf-8").splitlines()
    index = 0
    paragraph_lines: list[str] = []
    in_references = False

    def flush_paragraph() -> None:
        nonlocal paragraph_lines
        if paragraph_lines:
            add_paragraph(document, " ".join(line.strip() for line in paragraph_lines))
            paragraph_lines = []

    while index < len(lines):
        line = lines[index]
        stripped = line.strip()
        if not stripped:
            flush_paragraph()
            index += 1
            continue
        if stripped == "---":
            flush_paragraph()
            index += 1
            continue
        if stripped.startswith("```"):
            flush_paragraph()
            code_lines = []
            index += 1
            while index < len(lines) and not lines[index].strip().startswith("```"):
                code_lines.append(lines[index])
                index += 1
            index += 1
            paragraph = document.add_paragraph(style="Code Block")
            paragraph.add_run("\n".join(code_lines))
            continue
        if stripped.startswith("!["):
            flush_paragraph()
            add_image(document, stripped)
            index += 1
            continue
        if stripped.startswith("#"):
            flush_paragraph()
            match = re.match(r"^(#{1,4})\s+(.*)$", stripped)
            if match is None:
                raise ValueError(stripped)
            add_heading(document, match.group(2), len(match.group(1)))
            in_references = match.group(2) == "References"
            index += 1
            continue
        if stripped.startswith("|"):
            flush_paragraph()
            table_lines = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index])
                index += 1
            add_table(document, table_lines)
            continue
        if re.match(r"^[-*]\s+", stripped):
            flush_paragraph()
            paragraph = document.add_paragraph(style="List Bullet")
            add_inline(paragraph, re.sub(r"^[-*]\s+", "", stripped))
            index += 1
            continue
        if re.match(r"^\d+\.\s+", stripped):
            flush_paragraph()
            paragraph = document.add_paragraph(style="List Number")
            add_inline(paragraph, re.sub(r"^\d+\.\s+", "", stripped))
            if in_references:
                paragraph.paragraph_format.line_spacing = 1.0
                paragraph.paragraph_format.space_after = Pt(1)
                for run in paragraph.runs:
                    run.font.size = Pt(8.2)
            index += 1
            continue
        paragraph_lines.append(stripped)
        index += 1
    flush_paragraph()

    document.core_properties.title = lines[0].removeprefix("# ")
    document.core_properties.subject = "Scientific Data Data Descriptor"
    document.core_properties.author = "Jianxiang Huang; Xin Qiao; Shaoyong Lu"
    document.core_properties.keywords = (
        "GPCR; G protein; molecular dynamics; reduced trajectories; data descriptor"
    )
    document.core_properties.comments = (
        "v12 generated from auditable Markdown; release-dependent statements remain gated."
    )
    document.core_properties.modified = datetime.now(timezone.utc)
    target.parent.mkdir(parents=True, exist_ok=True)
    document.save(target)


def main() -> int:
    for source, target, line_numbers in BUILDS:
        render(source, target, line_numbers)
        print(f"Wrote {target.name}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
